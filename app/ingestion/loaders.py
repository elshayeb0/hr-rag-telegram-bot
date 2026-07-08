from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.ingestion.metadata import build_document_metadata


RawDocument = dict[str, Any]


def load_pdf(file_path: Path) -> list[RawDocument]:
    reader = PdfReader(str(file_path))
    documents: list[RawDocument] = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()

        if not text:
            continue

        documents.append(
            {
                "text": text,
                "source": file_path.name,
                "source_path": file_path,
                "file_type": "pdf",
                "page": page_index,
                "section_heading": None,
                "metadata": build_document_metadata(file_path),
            }
        )

    return documents


def _is_heading(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    return style_name.startswith("heading") or style_name in {"title", "subtitle"}


def load_docx(file_path: Path) -> list[RawDocument]:
    doc = DocxDocument(str(file_path))
    base_metadata = build_document_metadata(file_path)

    sections: list[RawDocument] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush_section() -> None:
        nonlocal current_lines, current_heading

        text = "\n".join(current_lines).strip()
        if not text:
            return

        sections.append(
            {
                "text": text,
                "source": file_path.name,
                "source_path": file_path,
                "file_type": "docx",
                "page": None,
                "section_heading": current_heading,
                "metadata": {
                    **base_metadata,
                    "section_heading": current_heading,
                },
            }
        )

        current_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        if _is_heading(paragraph):
            flush_section()
            current_heading = text
            current_lines = [text]
        else:
            current_lines.append(text)

    flush_section()

    if sections:
        return sections

    fallback_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip()).strip()

    if not fallback_text:
        return []

    return [
        {
            "text": fallback_text,
            "source": file_path.name,
            "source_path": file_path,
            "file_type": "docx",
            "page": None,
            "section_heading": None,
            "metadata": base_metadata,
        }
    ]


def load_xlsx(file_path: Path) -> list[RawDocument]:
    excel_file = pd.ExcelFile(file_path)
    documents: list[RawDocument] = []
    base_metadata = build_document_metadata(file_path)

    for sheet_name in excel_file.sheet_names:
        df = excel_file.parse(sheet_name)
        df = df.dropna(how="all")

        if df.empty:
            continue

        text = df.to_markdown(index=False)

        documents.append(
            {
                "text": text,
                "source": file_path.name,
                "source_path": file_path,
                "file_type": "xlsx",
                "sheet_name": sheet_name,
                "row_range": f"1-{len(df)}",
                "section_heading": sheet_name,
                "metadata": {
                    **base_metadata,
                    "section_heading": sheet_name,
                },
            }
        )

    return documents


def load_file(file_path: Path) -> list[RawDocument]:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(file_path)

    if suffix == ".docx":
        return load_docx(file_path)

    if suffix == ".xlsx":
        return load_xlsx(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


def load_directory(directory: Path) -> list[RawDocument]:
    documents: list[RawDocument] = []

    for file_path in sorted(directory.iterdir()):
        if file_path.is_dir() or file_path.name.startswith("."):
            continue

        documents.extend(load_file(file_path))

    return documents
